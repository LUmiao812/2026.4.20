
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy.signal import savgol_filter
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Inches
from io import BytesIO

# 设置图片清晰度
plt.rcParams['figure.dpi'] = 300

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['Microsoft YaHei']
plt.rcParams['axes.unicode_minus'] = False
from pathlib import Path
import pandas as pd

# 正确构建文件路径（不受转义影响）
# 注意这里用 r'...' 或者 Path 对象
file_path = Path(r"G:\10.11-10.15ZD 15 013579\ZZmerge.xlsx")

# 读取 Excel 文件
# 如果文件有中文路径，pandas 会自动支持0
df = pd.read_excel(file_path, header=None)

# 提取波长值（第一行）
wavelengths = df.iloc[0, 1:]

# 提取样本序号（第一列，从第二行开始）
sample_indices = df.iloc[1:, 0]

# 提取样本数据（去除第一行和第一列）
samples = df.iloc[1:, 1:]

# 创建一个新的 Word 文档
doc = Document()


def save_plot_to_doc(fig, title):
    img_stream = BytesIO()
    fig.savefig(img_stream, format='png')
    doc.add_heading(title, level=1)
    doc.add_picture(img_stream, width=Inches(6))
    img_stream.close()
    plt.close(fig)


# 绘制原始光谱图
fig = plt.figure(figsize=(12, 6))
for i in range(samples.shape[0]):
    plt.plot(wavelengths, samples.iloc[i, :], label=f'Sample {i + 1}')
plt.xlabel('Wavelength (nm)')
plt.ylabel('Reflectance')
plt.title('Raw spectral curve')
save_plot_to_doc(fig, 'Raw spectral curve')


# 执行 MSC 处理
def msc(spectra):
    mean_spectrum = np.mean(spectra, axis=0)
    msc_spectra = np.zeros_like(spectra)
    for i in range(spectra.shape[0]):
        fit = np.polyfit(mean_spectrum, spectra[i, :], 1)
        msc_spectra[i, :] = (spectra[i, :] - fit[1]) / fit[0]
    return msc_spectra


msc_spectra = msc(samples.values)
msc_df = pd.DataFrame(msc_spectra, columns=wavelengths)
# 添加样本序号列
msc_df.insert(0, 945, sample_indices.values)

# 绘制 MSC 处理后的光谱图
fig = plt.figure(figsize=(12, 6))
for i in range(msc_spectra.shape[0]):
    plt.plot(wavelengths, msc_spectra[i, :], label=f'Sample {i + 1}')
plt.xlabel('Wavelength  (nm)')
plt.ylabel('Reflectance')
plt.title('MSC preprocess spectral curve ')
save_plot_to_doc(fig, 'MSC preprocess spectral curve')

# 执行 SNV 处理
scaler = StandardScaler()
snv_spectra = scaler.fit_transform(samples.values)
snv_df = pd.DataFrame(snv_spectra, columns=wavelengths)
# 添加样本序号列
snv_df.insert(0, 945, sample_indices.values)

# 绘制 SNV 处理后的光谱图
fig = plt.figure(figsize=(12, 6))
for i in range(snv_spectra.shape[0]):
    plt.plot(wavelengths, snv_spectra[i, :], label=f'Sample {i + 1}')
plt.xlabel('Wavelength  (nm)')
plt.ylabel('Reflectance')
plt.title('SNV preprocess spectral curve')
save_plot_to_doc(fig, 'SNV preprocess spectral curve')

# 执行一阶导数处理
derivative_spectra = savgol_filter(samples.values, window_length=5, polyorder=2, deriv=1)
derivative_df = pd.DataFrame(derivative_spectra, columns=wavelengths)
# 添加样本序号列
derivative_df.insert(0, 945, sample_indices.values)

# 绘制一阶导数处理后的光谱图
fig = plt.figure(figsize=(12, 6))
for i in range(derivative_spectra.shape[0]):
    plt.plot(wavelengths, derivative_spectra[i, :], label=f'Sample {i + 1}')
plt.xlabel('Wavelength  (nm)')
plt.ylabel('Reflectance')
plt.title('the first derivate spectrometry preprocess spectral curve')
save_plot_to_doc(fig, '一阶导数处理后的光谱')

# 将处理后的数据保存到 Excel 文件
with pd.ExcelWriter('高光谱数据的处理结果.xlsx') as writer:
    msc_df.to_excel(writer, sheet_name='MSC 处理后', index=False)
    snv_df.to_excel(writer, sheet_name='SNV 处理后', index=False)
    derivative_df.to_excel(writer, sheet_name='一阶导数处理后', index=False)
print('Excel文件保存完成')
# 保存 Word 文档
doc.save('高光谱的处理结果报告.docx')
