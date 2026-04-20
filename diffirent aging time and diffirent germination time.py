import pandas as pd
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -------------------------- 1. 基础配置 --------------------------
# 解决中文显示问题
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False
# 图表样式配置
plt.rcParams['lines.linewidth'] = 2
plt.rcParams['axes.grid'] = True
plt.rcParams['grid.linestyle'] = '--'
plt.rcParams['grid.alpha'] = 0.6

# 文件路径配置
excel_path = r"G:\10.11-10.15ZD 15 013579\总光谱分类平均数据.xlsx"
# 图表保存文件夹（自动创建）
img_folder = r"G:\10.11-10.15ZD 15 013579\光谱图表"
# Word报告保存路径
word_path = r"G:\10.11-10.15ZD 15 013579\光谱分类平均数据报告.docx"

# 创建图表保存文件夹
if not os.path.exists(img_folder):
    os.makedirs(img_folder)

# -------------------------- 2. 读取Excel并绘制图表 --------------------------
# 获取所有工作表名称
excel_file = pd.ExcelFile(excel_path)
sheet_names = excel_file.sheet_names
# 存储图表路径，用于后续插入Word
img_paths = []

for sheet_name in sheet_names:
    # 读取工作表数据：第一列是标签，第一行是波长
    df = pd.read_excel(excel_path, sheet_name=sheet_name, index_col=0)
    # 提取波长（横坐标）、反射率（纵坐标）、曲线标签
    wavelength = df.columns.astype(float)
    reflectance = df.values
    labels = df.index.tolist()

    # 创建画布（预留图例空间）
    fig, ax = plt.subplots(figsize=(12, 7))
    # 绘制每条曲线
    for i, label in enumerate(labels):
        ax.plot(wavelength, reflectance[i], label=label, marker='', markersize=3)

    # 图表标题和坐标轴标签
    ax.set_title(f'{sheet_name} 光谱反射率曲线', fontsize=16, fontweight='bold', pad=20)
    ax.set_xlabel('Wavelength(nm)', fontsize=14, fontweight='bold')
    ax.set_ylabel('Reflectance', fontsize=14, fontweight='bold')
    ax.tick_params(axis='both', labelsize=12)

    # 图例设置：修改为内部左侧
    ax.legend(
        loc='upper left',  # 图例位置：图表内部左上角
        ncol=1,  # 纵向单列分布
        fontsize=10,
        frameon=True,
        shadow=True,
        bbox_to_anchor=(0.02, 0.98)  # 微调图例在左上角的位置（可选）
    )
    # 移除右侧预留空间的调整（因为图例在内部了）
    # plt.subplots_adjust(right=0.85)  # 注释掉这行

    # 保存图表
    img_name = f'{sheet_name}_光谱曲线.png'
    img_full_path = os.path.join(img_folder, img_name)
    plt.savefig(img_full_path, dpi=300, bbox_inches='tight')
    img_paths.append((sheet_name, img_full_path))  # 记录工作表名和图表路径

    # 关闭画布释放内存
    plt.close()
    print(f'✅ {sheet_name} 图表已保存：{img_full_path}')

# -------------------------- 3. 生成Word报告 --------------------------
# 创建新的Word文档
doc = Document()

# 1. 添加报告标题
title = doc.add_heading('光谱分类平均数据报告', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 标题居中

# 2. 添加报告说明（同步更新图例位置描述）
intro = doc.add_paragraph()
intro.add_run(
    '本报告包含6个工作表的光谱反射率曲线，横坐标为波长（nm），纵坐标为反射率（reflectance），所有曲线标签图例纵向分布在图表内部左侧。').bold = True

# 3. 逐个插入图表和对应的标题
for sheet_name, img_path in img_paths:
    # 添加工作表标题
    doc.add_heading(f'{sheet_name} 光谱反射率曲线', level=2)
    # 添加图表（设置宽度为6英寸，自适应Word页面）
    img_para = doc.add_paragraph()
    img_run = img_para.add_run()
    img_run.add_picture(img_path, width=Inches(6))
    # 图表居中
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 添加换行，分隔不同图表
    doc.add_paragraph('\n')

# 4. 保存Word文档
doc.save(word_path)

print(f'\n🎉 所有操作完成！')
print(f'📊 图表保存路径：{img_folder}')
print(f'📄 Word报告路径：{word_path}')