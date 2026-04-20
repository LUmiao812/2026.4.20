# 先统一导入所有需要的库（必须放在代码最开头）
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib as mpl
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
import warnings

# 屏蔽无关警告，保持控制台整洁
warnings.filterwarnings('ignore')

# ===================== 1. 配置Matplotlib中文字体+全局样式（兼容低版本） =====================
# 全局字体配置（支持中文）
mpl.rcParams["font.family"] = ["SimHei", "Microsoft YaHei"]
mpl.rcParams['axes.unicode_minus'] = False
# 全局样式优化（移除低版本不支持的参数）
mpl.rcParams['axes.grid'] = False  # 核心修改：先全局关闭网格，避免默认显示纵向网格
mpl.rcParams['xtick.labelsize'] = 10
mpl.rcParams['ytick.labelsize'] = 10
mpl.rcParams['axes.titlepad'] = 20
mpl.rcParams['figure.facecolor'] = 'white'
# 关闭柱子内部的网格线显示
mpl.rcParams['hatch.linewidth'] = 0.0

# ===================== 2. 读取数据并生成柱状图 =====================
# 读取Excel数据
df = pd.read_excel(r"G:\10.11-10.15ZD 15 013579\作图\八大模型准确率作图数据.xlsx")
models = df["模型"].tolist()
accuracies = df["准确度 % (验证)"].tolist()

# 确定最高准确率的模型
max_acc_idx = accuracies.index(max(accuracies))
max_acc_model = models[max_acc_idx]
max_acc_value = accuracies[max_acc_idx]

# ===================== 核心修改：替换为指定配色 =====================
# 普通模型：钢蓝色（#4682B4），最优模型：纯红色（#FF0000）
colors = ['#FF0000' if i == max_acc_idx else '#4682B4' for i in range(len(models))]

# 生成柱状图（单独一行，语法正确）
plt.figure(figsize=(8, 6))  # 核心修改：画幅调整为8英寸×6英寸
bars = plt.bar(
    models,
    accuracies,
    color=colors,  # 使用新的配色列表
    width=0.6,
    alpha=0.8,
    edgecolor='none',  # 彻底移除柱子线条
    linewidth=0,
    zorder=2           # 让柱子显示在网格线上方
)

# 添加数值标签
for bar, val in zip(bars, accuracies):
    height = bar.get_height()
    plt.text(
        bar.get_x() + bar.get_width()/2.,
        height + 0.5,
        f'{val:.4f}',
        ha='center',
        va='bottom',
        fontsize=9,
        fontweight='bold',
        zorder=3  # 标签显示在最上层
    )

# 图表样式设置
plt.title("Comparison of model accuracy rates", fontsize=14, fontweight='bold', pad=20)
plt.ylabel("Accuracy", fontsize=11)
plt.xlabel("Model", fontsize=11)
plt.ylim(0, max(accuracies) + 5)
plt.xticks(rotation=45, ha="right", fontsize=10)
# 核心修改：仅开启y轴（横向）网格，完全关闭x轴（纵向）网格
plt.grid(alpha=0.3, axis='y', zorder=1)  # 只显示横向网格，无纵向网格
plt.tight_layout(pad=2)

# 保存图片
img_path = "模型准确率对比图.png"
plt.savefig(
    img_path,
    dpi=300,
    bbox_inches='tight',
    facecolor='white'
)
plt.close()

# ===================== 3. 生成Word报告 =====================
doc = Document()

# 3.1 设置Word字体函数
def set_doc_font(paragraph, font_name="微软雅黑", font_size=12):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# 3.2 添加标题
title_para = doc.add_heading("八大模型准确率分析报告", level=0)
title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
set_doc_font(title_para, font_size=16)

# 3.3 数据概述
intro_para = doc.add_paragraph()
intro_para.add_run("一、数据概述\n").font.bold = True
set_doc_font(intro_para, font_size=14)
intro_content = doc.add_paragraph(
    f"本次分析共涉及{len(models)}个模型的准确率对比，数据来源于Excel文件：\n"
    f"G:\\10.11-10.15ZD 15 013579\\作图\\八大模型准确率作图数据.xlsx。\n\n"
    f"数据维度说明：\n"
    f"• 模型数量：{len(models)}个\n"
    f"• 评估指标：准确度 % (验证)\n"
    f"• 最优模型：【{max_acc_model}】，准确率为{max_acc_value:.4f}（{max_acc_value:.2f}%）。"
)
set_doc_font(intro_content)

# 3.4 准确率可视化
chart_para = doc.add_paragraph()
chart_para.add_run("二、准确率可视化\n").font.bold = True
set_doc_font(chart_para, font_size=14)
img_para = doc.add_paragraph()
img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
img_run = img_para.add_run()
img_run.add_picture(img_path, width=Inches(7))

# 3.5 模型性能排序表格
rank_para = doc.add_paragraph()
rank_para.add_run("三、模型性能排序\n").font.bold = True
set_doc_font(rank_para, font_size=14)

# 生成排序数据
model_rank = pd.DataFrame({
    "模型名称": models,
    "准确率": accuracies
}).sort_values(by="准确率", ascending=False).reset_index(drop=True)
model_rank["排名"] = range(1, len(model_rank)+1)
model_rank = model_rank[["排名", "模型名称", "准确率"]]

# 创建Word表格
table = doc.add_table(rows=len(model_rank)+1, cols=3)
table.style = 'Table Grid'
# 表头
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '排名'
hdr_cells[1].text = '模型名称'
hdr_cells[2].text = '准确率'
set_doc_font(table.rows[0].cells[0].paragraphs[0], font_size=11)
set_doc_font(table.rows[0].cells[1].paragraphs[0], font_size=11)
set_doc_font(table.rows[0].cells[2].paragraphs[0], font_size=11)

# 填充表格数据
for idx, row in model_rank.iterrows():
    row_cells = table.rows[idx+1].cells
    row_cells[0].text = str(row["排名"])
    row_cells[1].text = row["模型名称"]
    row_cells[2].text = f"{row['准确率']:.4f}"
    set_doc_font(row_cells[0].paragraphs[0], font_size=10)
    set_doc_font(row_cells[1].paragraphs[0], font_size=10)
    set_doc_font(row_cells[2].paragraphs[0], font_size=10)

# 3.6 关键结论
conclusion_para = doc.add_paragraph()
conclusion_para.add_run("四、关键结论\n").font.bold = True
set_doc_font(conclusion_para, font_size=14)
conclusion_content = doc.add_paragraph(
    f"1. 性能最优模型：{max_acc_model}，准确率达到{max_acc_value:.4f}（{max_acc_value:.2f}%），在所有模型中表现最佳；\n"
    f"2. 模型性能分布：本次分析的{len(models)}个模型中，准确率差异主要集中在{min(accuracies):.4f}-{max(accuracies):.4f}区间；\n"
    f"3. 选型建议：优先考虑{max_acc_model}作为核心模型，其余模型可根据业务场景（如计算效率、解释性）作为补充验证；\n"
    f"4. 数据可靠性：所有准确率数据均基于验证集计算，具备较好的泛化性参考价值。"
)
set_doc_font(conclusion_content)

# 3.7 保存Word文档
doc.save(r"G:\10.11-10.15ZD 15 013579\八大模型准确率分析报告.docx")
print("✅ Word报告生成完成！")
print(f"📄 报告路径：G:\\10.11-10.15ZD 15 013579\\八大模型准确率分析报告.docx")
print(f"📊 图表路径：{img_path}")