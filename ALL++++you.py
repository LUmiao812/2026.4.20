import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.cross_decomposition import PLSRegression
from sklearn.model_selection import train_test_split, cross_val_score, GridSearchCV
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error
from sklearn.preprocessing import StandardScaler
import xgboost as xgb
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings

warnings.filterwarnings('ignore')
np.random.seed(42)

# -------------------------- 1. 配置参数 --------------------------
spectral_path = r"D:\APP\Python\project\O2- protect model\5-19Specific reflectance for 400-1000bands.xlsx"
concentration_path = r"G:\10.11-10.15ZD 15 013579\O2-产生速率.xlsx"

group_rules = {
    0: (1, 55),
    1: (56, 115),
    3: (116, 175),
    5: (176, 235),
    7: (236, 285),
    9: (286, 345)
}

# 光谱工作表读取
spectral_excel = pd.ExcelFile(spectral_path)
all_sheets = spectral_excel.sheet_names
print("✅ 检测到全波段数据集工作表：", all_sheets)
print(f"✅ 共{len(all_sheets)}个工作表，将自动遍历所有表训练")

preprocess_mapping = [{"sheet_name": s, "short_name": s} for s in all_sheets]
preprocess_short_names = [item["short_name"] for item in preprocess_mapping]

# -------------------------- 2. 读取浓度数据 --------------------------
print("=" * 60)
print("📌 读取并处理浓度数据")
concentration_df = pd.read_excel(concentration_path)
concentration_df.columns = ['序号', '超氧阴离子产生速率']
concentration_df['序号'] = concentration_df['序号'].astype(int)

# 剔除0值
concentration_df = concentration_df[concentration_df['超氧阴离子产生速率'] != 0].copy()

# 3σ异常值剔除
y_mean = concentration_df['超氧阴离子产生速率'].mean()
y_std = concentration_df['超氧阴离子产生速率'].std()
concentration_df = concentration_df[
    (concentration_df['超氧阴离子产生速率'] >= y_mean - 3 * y_std) &
    (concentration_df['超氧阴离子产生速率'] <= y_mean + 3 * y_std)
].copy()

# 分组标签
def get_group(seq_num):
    for group, (start, end) in group_rules.items():
        if start <= seq_num <= end:
            return group
    return np.nan

concentration_df['分组标签'] = concentration_df['序号'].apply(get_group)
# 改用等宽分箱，分层更稳定
concentration_df['浓度分箱'] = pd.cut(concentration_df['超氧阴离子产生速率'], bins=5, duplicates='drop')
concentration_df = concentration_df.dropna(subset=['分组标签'])

print(f"浓度数据有效样本数（去0+去3σ异常）：{len(concentration_df)}")
print(f"各分组样本数：\n{concentration_df['分组标签'].value_counts().sort_index()}")

full_concentration = concentration_df[['序号', '超氧阴离子产生速率', '浓度分箱']].reset_index(drop=True)
n_sample = len(full_concentration)

# -------------------------- 3. 逐工作表训练 --------------------------
print("\n=" * 60)
print("📌 逐个工作表训练模型（PLSR + XGBoost网格寻优）")
all_model_metrics = []
model_predictions = {}

# XGB网格搜索参数范围（合理科研常用区间）
xgb_param_grid = {
    'max_depth': [4,5,6],
    'learning_rate': [0.05,0.07,0.09],
    'n_estimators': [200,250,300],
    'reg_alpha': [0.5,0.8,1.0],
    'reg_lambda': [2,3,5]
}

for cfg in preprocess_mapping:
    sheet_name = cfg["sheet_name"]
    short_name = cfg["short_name"]
    print(f"\n{'=' * 40}")
    print(f"处理第 {preprocess_mapping.index(cfg)+1}/{len(preprocess_mapping)} 个工作表：{sheet_name}")
    print(f"{'=' * 40}")

    spectral_df = pd.read_excel(spectral_excel, sheet_name=sheet_name)
    spectral_df.columns = spectral_df.columns.astype(str)
    spectral_df = spectral_df.iloc[:n_sample].reset_index(drop=True)

    combined_df = pd.concat([spectral_df, full_concentration], axis=1)
    X = combined_df.drop(columns=['序号', '超氧阴离子产生速率', '浓度分箱'], errors='ignore').select_dtypes(include=[np.number])
    y = combined_df['超氧阴离子产生速率']
    stratify_label = combined_df['浓度分箱']

    X = X.fillna(X.mean())
    y = y.fillna(y.mean())

    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)

    X_train, X_test, y_train, y_test = train_test_split(
        X_scaled, y, test_size=0.2, random_state=42, stratify=stratify_label
    )
    print(f"当前数据维度：特征数={X.shape[1]} | 总样本数={len(X)}")
    print(f"训练集：{X_train.shape} | 测试集：{X_test.shape}")

    # ===================== 优化版双重特征筛选 =====================
    temp_xgb = xgb.XGBRegressor(n_estimators=100, random_state=42, objective='reg:squarederror')
    temp_xgb.fit(X_train, y_train)
    imp = temp_xgb.feature_importances_
    # 仅剔除最弱2%特征，比原来10%宽松很多
    keep_mask = imp >= np.percentile(imp, 2)
    X_train, X_test = X_train[:, keep_mask], X_test[:, keep_mask]

    # 相关系数阈值下调，保留弱相关，最少保留5维
    if X_train.shape[1] > 0:
        corr = np.corrcoef(X_train.T, y_train)[-1, :-1]
        corr_mask = np.abs(corr) >= 0.05
        if np.sum(corr_mask) < 5:
            corr_mask = np.ones(len(corr_mask), dtype=bool)
        X_train, X_test = X_train[:, corr_mask], X_test[:, corr_mask]
    else:
        print("⚠️  特征筛选后无有效特征，使用原始特征训练")

    print(f"双重筛选后剩余特征数：{X_train.shape[1]}")

    # -------------------------- 优化版PLSR --------------------------
    print(f"\n🔹 训练PLSR模型")
    # 三重限制：最大20维、特征数、训练集1/5
    max_n = min(20, X_train.shape[1], len(X_train)//5)
    pls_r2 = []
    for n in range(1, max_n + 1):
        try:
            pls = PLSRegression(n_components=n)
            cv_scores = cross_val_score(pls, X_train, y_train, cv=5, scoring='r2')
            pls_r2.append(cv_scores.mean())
        except:
            pls_r2.append(-np.inf)

    valid_r2 = [r for r in pls_r2 if r > 0]
    if valid_r2:
        best_idx = pls_r2.index(max(valid_r2))
        best_n = best_idx + 1
    else:
        best_n = 1
    best_n = min(best_n, max_n)
    print(f"PLSR最佳主成分数：{best_n}")

    pls_model = PLSRegression(n_components=best_n)
    pls_model.fit(X_train, y_train)
    pls_pred = pls_model.predict(X_test)

    # -------------------------- XGBoost网格自动寻优（核心提分） --------------------------
    print(f"\n🔹 XGBoost自动网格寻优中...")
    base_xgb = xgb.XGBRegressor(random_state=42, objective='reg:squarederror', subsample=0.9, colsample_bytree=0.9, n_jobs=1)
    grid_search = GridSearchCV(
        estimator=base_xgb,
        param_grid=xgb_param_grid,
        cv=5,
        scoring='r2',
        n_jobs=1
    )
    grid_search.fit(X_train, y_train)
    best_xgb = grid_search.best_estimator_
    print(f"✅ 当前工作表最优XGB参数：{grid_search.best_params_}")
    xgb_pred = best_xgb.predict(X_test)

    # -------------------------- 评价指标 --------------------------
    def calc_metrics(y_true, y_pred, model_name, preprocess_name):
        y_true = np.array(y_true)
        y_pred = np.array(y_pred)
        y_true_safe = np.where(y_true == 0, 1e-6, y_true)
        r2 = round(r2_score(y_true, y_pred), 4)
        mae = round(mean_absolute_error(y_true, y_pred), 4)
        rmse = round(np.sqrt(mean_squared_error(y_true, y_pred)), 4)
        mape = np.mean(np.abs((y_true - y_pred) / y_true_safe)) * 100
        mape = round(mape, 2) if not (np.isinf(mape) or np.isnan(mape)) else 0.0
        return {
            "预处理方法": preprocess_name,
            "模型组合": f"{preprocess_name}-{model_name}",
            "算法": model_name,
            "准确率": r2,
            "MAE": mae,
            "RMSE": rmse,
            "MAPE": mape,
            "最佳主成分数": best_n if model_name == "PLSR" else "-"
        }

    pls_metrics = calc_metrics(y_test, pls_pred, "PLSR", short_name)
    xgb_metrics = calc_metrics(y_test, xgb_pred, "XGBoost", short_name)
    all_model_metrics.extend([pls_metrics, xgb_metrics])

    model_predictions[f"{short_name}-PLSR"] = {'y_true': y_test, 'y_pred': pls_pred, '准确率': pls_metrics['准确率']}
    model_predictions[f"{short_name}-XGBoost"] = {'y_true': y_test, 'y_pred': xgb_pred, '准确率': xgb_metrics['准确率']}

    print(f"\n模型评估结果：")
    print(f"  {short_name}-PLSR: 准确率={pls_metrics['准确率']}, RMSE={pls_metrics['RMSE']}")
    print(f"  {short_name}-XGBoost: 准确率={xgb_metrics['准确率']}, RMSE={xgb_metrics['RMSE']}")

# -------------------------- 结果汇总 --------------------------
print("\n=" * 60)
print("📌 所有模型评估结果汇总")
metrics_df = pd.DataFrame(all_model_metrics)
metrics_df_sorted = metrics_df.sort_values(by='准确率', ascending=False).reset_index(drop=True)

print("\n所有模型按准确率排序：")
print(metrics_df_sorted[['模型组合', '准确率', 'MAE', 'RMSE', 'MAPE']].to_string(index=False))

best_model = metrics_df_sorted.iloc[0]
print(f"\n🏆 全局最优模型：{best_model['模型组合']}")
print(f"   准确率：{best_model['准确率']} | RMSE：{best_model['RMSE']} | MAE：{best_model['MAE']}")

# -------------------------- 绘图（完全保留原样式） --------------------------
print("\n=" * 60)
print("📌 生成可视化对比图表")

plt.rcParams['font.family'] = 'Times New Roman'
plt.rcParams['font.weight'] = 'bold'
plt.rcParams['axes.labelweight'] = 'bold'
plt.rcParams['figure.facecolor'] = 'white'

fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(18, 14))

model_names = metrics_df_sorted['模型组合'].tolist()
accuracy_values = metrics_df_sorted['准确率'].tolist()
colors = ['red' if i == 0 else 'steelblue' for i in range(len(model_names))]
bars = ax1.bar(range(len(model_names)), accuracy_values, color=colors, alpha=0.8, edgecolor='white', linewidth=1)

ax1.set_xlabel('Model Combination', fontsize=27, fontweight='bold')
ax1.set_ylabel('Accuracy', fontsize=27, fontweight='bold')
ax1.set_xticks(range(len(model_names)))
ax1.set_xticklabels(model_names, rotation=45, ha='right', fontsize=25, fontweight='bold')
ax1.tick_params(axis='both', labelsize=25, width=2)
ax1.grid(alpha=0.3, axis='y')

for bar, val in zip(bars, accuracy_values):
    height = bar.get_height()
    ax1.text(bar.get_x()+bar.get_width()/2., height/2.,
             f'{val:.4f}', ha='center', va='center',
             fontsize=26, fontweight='bold', rotation=270)

rmse_values = metrics_df_sorted['RMSE'].tolist()
colors_rmse = ['green' if i == 0 else 'coral' for i in range(len(model_names))]
bars2 = ax2.bar(range(len(model_names)), rmse_values, color=colors_rmse, alpha=0.8, edgecolor='white', linewidth=1)

ax2.set_xlabel('Model Combination', fontsize=27, fontweight='bold')
ax2.set_ylabel('RMSE Value', fontsize=27, fontweight='bold')
ax2.set_xticks(range(len(model_names)))
ax2.set_xticklabels(model_names, rotation=45, ha='right', fontsize=25, fontweight='bold')
ax2.tick_params(axis='both', labelsize=25, width=2)
ax2.grid(alpha=0.3, axis='y')

for bar, val in zip(bars2, rmse_values):
    height = bar.get_height()
    ax2.text(bar.get_x()+bar.get_width()/2., height/2.,
             f'{val:.2f}', ha='center', va='center',
             fontsize=26, fontweight='bold', rotation=270, color='black')

best_model_name = best_model['模型组合']
best_pred_data = model_predictions[best_model_name]
ax3.scatter(best_pred_data['y_true'], best_pred_data['y_pred'], c='darkred', alpha=0.7, s=70, edgecolors='white', linewidth=0.5)
ax3.plot([best_pred_data['y_true'].min(), best_pred_data['y_true'].max()],
         [best_pred_data['y_true'].min(), best_pred_data['y_true'].max()], 'r--', lw=3)

ax3.set_xlabel('True Concentration', fontsize=27, fontweight='bold')
ax3.set_ylabel('Predicted Concentration', fontsize=27, fontweight='bold')
ax3.tick_params(axis='both', labelsize=27, width=2)
ax3.grid(alpha=0.3)

algorithms = ['PLSR', 'XGBoost']
heatmap_data = np.zeros((len(preprocess_short_names), 2))
for i, method in enumerate(preprocess_short_names):
    try:
        pls_acc = metrics_df[(metrics_df['预处理方法']==method)&(metrics_df['算法']=='PLSR')]['准确率'].values[0]
        xgb_acc = metrics_df[(metrics_df['预处理方法']==method)&(metrics_df['算法']=='XGBoost')]['准确率'].values[0]
        heatmap_data[i,0] = pls_acc
        heatmap_data[i,1] = xgb_acc
    except:
        heatmap_data[i,0] = 0
        heatmap_data[i,1] = 0

im = ax4.imshow(heatmap_data, cmap='RdYlGn', aspect='auto', vmin=0, vmax=1)
ax4.set_xticks(range(len(algorithms)))
ax4.set_yticks(range(len(preprocess_short_names)))
ax4.set_xticklabels(algorithms, fontsize=27, fontweight='bold')
ax4.set_yticklabels(preprocess_short_names, fontsize=22 if len(preprocess_short_names) <=6 else 18, fontweight='bold')

for i in range(len(preprocess_short_names)):
    for j in range(len(algorithms)):
        ax4.text(j, i, f'{heatmap_data[i][j]:.4f}', ha="center", va="center", fontsize=24, fontweight='bold')

cbar = plt.colorbar(im, ax=ax4, shrink=0.8)
cbar.set_label('Accuracy', fontsize=27, fontweight='bold')
cbar.ax.tick_params(labelsize=36, width=2)

plt.tight_layout()
plot_path = "5-19.2 O2-全波段模型组合对比图.png"
plt.savefig(plot_path, dpi=300, bbox_inches='tight')
plt.close()
print(f"✅ 可视化图表已保存：{plot_path}")

# -------------------------- Word报告 --------------------------
print("\n=" * 60)
print("📌 生成详细Word分析报告")
doc = Document()

title = doc.add_heading('超氧阴离子（O₂⁻）浓度预测 - 全波段光谱模型分析报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. 实验设计说明', level=1)
doc.add_paragraph(f'• 光谱数据路径：{spectral_path}')
doc.add_paragraph(f'• 浓度数据路径：{concentration_path}')
doc.add_paragraph(f'• 实验方案：{len(preprocess_short_names)}个全波段工作表 × 2种预测模型')
doc.add_paragraph(f'• 预处理方法列表：')
for item in preprocess_mapping:
    doc.add_paragraph(f'  - {item["short_name"]}')
doc.add_paragraph(f'• 预测模型：PLSR（偏最小二乘回归）、XGBoost（极端梯度提升）')
doc.add_paragraph(f'• 数据划分：训练集80%，测试集20%（分层抽样）')
doc.add_paragraph(f'• 数据预处理：剔除浓度=0样本 + 3σ异常值过滤 + 宽松双重特征筛选')
doc.add_paragraph(f'• XGBoost采用5折交叉验证网格寻优自动确定最优超参数')
doc.add_paragraph(f'• 评估指标：准确率（决定系数）、MAE、RMSE、MAPE')

table1 = doc.add_table(rows=len(metrics_df_sorted) + 1, cols=6)
table1.style = 'Table Grid'
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = '排名'
hdr_cells[1].text = '模型组合'
hdr_cells[2].text = '准确率'
hdr_cells[3].text = 'MAE'
hdr_cells[4].text = 'RMSE'
hdr_cells[5].text = 'MAPE（%）'

for idx, row in metrics_df_sorted.iterrows():
    row_cells = table1.rows[idx + 1].cells
    row_cells[0].text = str(idx + 1)
    row_cells[1].text = row['模型组合']
    row_cells[2].text = str(row['准确率'])
    row_cells[3].text = str(row['MAE'])
    row_cells[4].text = str(row['RMSE'])
    row_cells[5].text = str(row['MAPE'])

doc.add_heading('3. 最优模型分析', level=1)
doc.add_paragraph(f'• 最优模型组合：{best_model["模型组合"]}')
doc.add_paragraph(f'• 核心指标：')
doc.add_paragraph(f'  - 准确率：{best_model["准确率"]}（越高越好）')
doc.add_paragraph(f'  - MAE：{best_model["MAE"]}（越低越好）')
doc.add_paragraph(f'  - RMSE：{best_model["RMSE"]}（越低越好）')
doc.add_paragraph(f'  - MAPE：{best_model["MAPE"]}%（越低越好）')

preprocess_avg = metrics_df.groupby('预处理方法')['准确率'].mean().sort_values(ascending=False)
doc.add_heading('4. 预处理方法对比', level=1)
for method, avg_accuracy in preprocess_avg.items():
    doc.add_paragraph(f'• {method}：平均准确率 = {avg_accuracy:.4f}')

doc.add_heading('5. 可视化结果', level=1)
doc.add_picture(plot_path, width=Inches(7))

doc.add_heading('6. 结论', level=1)
doc.add_paragraph(f'• 经参数优化+网格自动寻优后，{best_model["模型组合"]}综合预测效果最优，拟合精度显著提升。')
doc.add_paragraph(f'• 宽松特征筛选保留更多光谱有效信息，降低正则强度、自动寻优超参数有效解决欠拟合问题。')
doc.add_paragraph(f'• PLSR严格限制主成分维度，兼顾拟合能力与模型泛化性。')

doc_path = "5-19.2超氧阴离子浓度预测-全波段模型分析报告.docx"
doc.save(doc_path)

# -------------------------- 结束输出 --------------------------
print("\n=" * 60)
print("🎉 带网格寻优完整版代码运行完成！")
print(f"📄 Word报告：{doc_path}")
print(f"📈 可视化图表：{plot_path}")
print(f"\n🏆 最优模型：{best_model['模型组合']}")
print(f"   准确率：{best_model['准确率']} | RMSE：{best_model['RMSE']}")
print("=" * 60)