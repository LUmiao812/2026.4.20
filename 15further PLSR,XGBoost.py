import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.cross_decomposition import PLSRegression
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error
from sklearn.preprocessing import StandardScaler
import xgboost as xgb
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings

warnings.filterwarnings('ignore')

# -------------------------- 1. 配置参数与数据分组规则 --------------------------
# 文件路径
spectral_path = r"D:\APP\Python\project\1.4 O2- 预处理后的CatBoost和GBDT 15个特征波长\1.4 0H TZspectral_model_results_with_weights_summary\预处理后15个特征波段的具体反射值.xlsx"
concentration_path = r"G:\10.11-10.15ZD 15 013579\O2-浓度.xlsx"

# 浓度数据分组规则（序号范围: 分组标签）
group_rules = {
    0: (1, 55),
    1: (56, 115),
    3: (116, 175),
    5: (176, 235),
    7: (236, 285),
    9: (286, 345)
}

# 预处理方法名称映射（严格匹配工作表名称）
preprocess_mapping = [
    {"sheet_name": "MSC 处理后CatBoost", "short_name": "MSC-CatBoost"},
    {"sheet_name": "MSC 处理后GBDT", "short_name": "MSC-GBDT"},
    {"sheet_name": "SNV 处理后CatBoost", "short_name": "SNV-CatBoost"},
    {"sheet_name": "SNV 处理后GBDT", "short_name": "SNV-GBDT"},
    {"sheet_name": "一阶导数处理后CatBoost", "short_name": "D1-CatBoost"},
    {"sheet_name": "一阶导数处理后GBDT", "short_name": "D1-GBDT"}
]

# 提取短名称列表
preprocess_short_names = [item["short_name"] for item in preprocess_mapping]

# -------------------------- 2. 读取浓度数据并标记分组 --------------------------
print("=" * 60)
print("📌 读取并处理浓度数据")
concentration_df = pd.read_excel(concentration_path)
# 统一列名并转为字符串
concentration_df.columns = ['序号', '超氧阴离子浓度']
concentration_df['序号'] = concentration_df['序号'].astype(int)


# 添加分组标签列
def get_group(seq_num):
    for group, (start, end) in group_rules.items():
        if start <= seq_num <= end:
            return group
    return np.nan


concentration_df['分组标签'] = concentration_df['序号'].apply(get_group)
# 过滤有效分组数据
concentration_df = concentration_df.dropna(subset=['分组标签'])
print(f"浓度数据有效样本数：{len(concentration_df)}")
print(f"各分组样本数：\n{concentration_df['分组标签'].value_counts().sort_index()}")

# 提取完整的浓度数据（所有分组合并）
full_concentration = concentration_df[['序号', '超氧阴离子浓度']].reset_index(drop=True)

# -------------------------- 3. 逐个处理工作表训练模型 --------------------------
print("\n=" * 60)
print("📌 逐个预处理方法训练模型（PLSR + XGBoost）")
# 存储所有模型的评估结果
all_model_metrics = []
# 存储每个模型的预测结果（用于可视化）
model_predictions = {}

# 读取光谱文件
spectral_excel = pd.ExcelFile(spectral_path)
all_sheet_names = spectral_excel.sheet_names
# 只处理前6个工作表
target_sheets = all_sheet_names[:6]

for sheet_idx, sheet_name in enumerate(target_sheets):
    print(f"\n{'=' * 40}")
    print(f"处理第{sheet_idx + 1}个预处理方法：{sheet_name}")
    print(f"{'=' * 40}")

    # 获取对应的短名称
    short_name = preprocess_mapping[sheet_idx]["short_name"]

    # 读取当前预处理方法的光谱数据
    spectral_df = pd.read_excel(spectral_excel, sheet_name=sheet_name)
    # 统一列名为字符串
    spectral_df.columns = spectral_df.columns.astype(str)
    # 只保留有效样本数
    spectral_df = spectral_df.iloc[:len(full_concentration)].reset_index(drop=True)

    # 合并光谱数据和浓度数据
    combined_df = pd.concat([
        spectral_df,
        full_concentration
    ], axis=1)

    # -------------------------- 3.1 数据预处理 --------------------------
    # 分离特征和目标变量
    X = combined_df.drop(columns=['序号', '超氧阴离子浓度']).select_dtypes(include=[np.number])
    y = combined_df['超氧阴离子浓度']

    # 确保特征列名全为字符串
    X.columns = X.columns.astype(str)
    # 缺失值处理
    X = X.fillna(X.mean())
    y = y.fillna(y.mean())

    # 标准化
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)

    # 划分训练集和测试集（8:2）
    X_train, X_test, y_train, y_test = train_test_split(
        X_scaled, y, test_size=0.2, random_state=42
    )
    print(f"当前预处理方法数据维度：")
    print(f"  特征数：{X.shape[1]} | 总样本数：{len(X)}")
    print(f"  训练集：{X_train.shape} | 测试集：{X_test.shape}")

    # -------------------------- 3.2 训练PLSR模型 --------------------------
    print(f"\n🔹 训练PLSR模型")
    # 自动选择最佳主成分数
    pls_r2 = []
    max_n = min(15, X_train.shape[1])  # 最多15个主成分
    for n in range(1, max_n + 1):
        pls = PLSRegression(n_components=n)
        cv_scores = cross_val_score(pls, X_train, y_train, cv=5, scoring='r2')
        pls_r2.append(cv_scores.mean())
    best_n = np.argmax(pls_r2) + 1
    print(f"  PLSR最佳主成分数：{best_n}（CV 准确率：{max(pls_r2):.4f}）")

    # 训练最终PLSR模型
    pls_model = PLSRegression(n_components=best_n)
    pls_model.fit(X_train, y_train)
    pls_pred = pls_model.predict(X_test)

    # -------------------------- 3.3 训练XGBoost模型 --------------------------
    print(f"\n🔹 训练XGBoost模型")
    xgb_model = xgb.XGBRegressor(
        n_estimators=150,
        learning_rate=0.08,
        max_depth=6,
        subsample=0.8,
        colsample_bytree=0.8,
        random_state=42,
        objective='reg:squarederror'
    )
    xgb_model.fit(X_train, y_train)
    xgb_pred = xgb_model.predict(X_test)


    # -------------------------- 3.4 计算评估指标 --------------------------
    def calc_metrics(y_true, y_pred, model_name, preprocess_name):
        # 处理MAPE计算中的除以0和inf问题
        y_true = np.array(y_true)
        y_pred = np.array(y_pred)

        # 替换0值为极小值，避免除以0
        y_true_safe = np.where(y_true == 0, 1e-6, y_true)

        # 计算指标
        r2 = round(r2_score(y_true, y_pred), 4)
        mae = round(mean_absolute_error(y_true, y_pred), 4)
        rmse = round(np.sqrt(mean_squared_error(y_true, y_pred)), 4)

        # 计算MAPE并处理inf
        mape = np.mean(np.abs((y_true - y_pred) / y_true_safe)) * 100
        if np.isinf(mape) or np.isnan(mape):
            mape = 0.0
        else:
            mape = round(mape, 2)

        return {
            "预处理方法": preprocess_name,
            "模型组合": f"{preprocess_name}-{model_name}",
            "算法": model_name,
            "准确率": r2,  # 字段名改为准确率
            "MAE": mae,
            "RMSE": rmse,
            "MAPE": mape,
            "最佳主成分数": best_n if model_name == "PLSR" else "-"
        }


    # 计算PLSR指标
    pls_metrics = calc_metrics(y_test, pls_pred, "PLSR", short_name)
    all_model_metrics.append(pls_metrics)
    # 存储PLSR预测结果
    model_predictions[f"{short_name}-PLSR"] = {
        'y_true': y_test,
        'y_pred': pls_pred,
        '准确率': pls_metrics['准确率']  # 字段名改为准确率
    }

    # 计算XGBoost指标
    xgb_metrics = calc_metrics(y_test, xgb_pred, "XGBoost", short_name)
    all_model_metrics.append(xgb_metrics)
    # 存储XGBoost预测结果
    model_predictions[f"{short_name}-XGBoost"] = {
        'y_true': y_test,
        'y_pred': xgb_pred,
        '准确率': xgb_metrics['准确率']  # 字段名改为准确率
    }

    # 打印当前预处理方法的模型结果
    print(f"\n当前预处理方法模型评估结果：")
    print(f"  {short_name}-PLSR: 准确率={pls_metrics['准确率']}, RMSE={pls_metrics['RMSE']}")
    print(f"  {short_name}-XGBoost: 准确率={xgb_metrics['准确率']}, RMSE={xgb_metrics['RMSE']}")

# -------------------------- 4. 整理并展示所有模型结果 --------------------------
print("\n=" * 60)
print("📌 所有模型评估结果汇总")
# 转换为DataFrame
metrics_df = pd.DataFrame(all_model_metrics)
# 按准确率降序排序（准确率从高到低）
metrics_df_sorted = metrics_df.sort_values(by='准确率', ascending=False).reset_index(drop=True)

# 打印完整结果
print("\n所有模型按准确率排序：")
print(metrics_df_sorted[['模型组合', '准确率', 'MAE', 'RMSE', 'MAPE']].to_string(index=False))

# 找出最优模型
best_model = metrics_df_sorted.iloc[0]
print(f"\n🏆 最优模型：{best_model['模型组合']}")
print(f"   准确率：{best_model['准确率']} | RMSE：{best_model['RMSE']} | MAE：{best_model['MAE']}")

# -------------------------- 5. 可视化对比（替换R²为准确率） --------------------------
print("\n=" * 60)
print("📌 生成模型对比可视化图表")
# 注意：这里不设置中文字体，因为图表将显示英文
plt.rcParams['figure.facecolor'] = 'white'

# 创建2行2列的可视化图表
fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(16, 12))

# 5.1 所有模型准确率对比（柱状图）
model_names = metrics_df_sorted['模型组合'].tolist()
accuracy_values = metrics_df_sorted['准确率'].tolist()  # 改为准确率

# 为最优模型标红
colors = ['red' if i == 0 else 'steelblue' for i in range(len(model_names))]
bars = ax1.bar(range(len(model_names)), accuracy_values, color=colors, alpha=0.8, edgecolor='white', linewidth=1)
ax1.set_xlabel('Model Combination', fontsize=11)
ax1.set_ylabel('Accuracy', fontsize=11)  # 替换R²值为准确率
ax1.set_title('Accuracy Comparison of All Models', fontsize=14, fontweight='bold')  # 替换R²为准确率
ax1.set_xticks(range(len(model_names)))
ax1.set_xticklabels(model_names, rotation=45, ha='right', fontsize=9)
ax1.grid(alpha=0.3, axis='y')

# 添加数值标签
for bar, val in zip(bars, accuracy_values):
    height = bar.get_height()
    ax1.text(bar.get_x() + bar.get_width() / 2., height,
             f'{val:.4f}', ha='center', va='bottom', fontsize=8)

# 5.2 所有模型RMSE对比
rmse_values = metrics_df_sorted['RMSE'].tolist()
colors_rmse = ['green' if i == 0 else 'coral' for i in range(len(model_names))]
bars2 = ax2.bar(range(len(model_names)), rmse_values, color=colors_rmse, alpha=0.8, edgecolor='white', linewidth=1)
ax2.set_xlabel('Model Combination', fontsize=11)
ax2.set_ylabel('RMSE Value', fontsize=11)
ax2.set_title('RMSE Comparison of All Models', fontsize=14, fontweight='bold')
ax2.set_xticks(range(len(model_names)))
ax2.set_xticklabels(model_names, rotation=45, ha='right', fontsize=9)
ax2.grid(alpha=0.3, axis='y')

# 添加数值标签
for bar, val in zip(bars2, rmse_values):
    height = bar.get_height()
    ax2.text(bar.get_x() + bar.get_width() / 2., height,
             f'{val:.4f}', ha='center', va='bottom', fontsize=8)

# 5.3 最优模型预测散点图
best_model_name = best_model['模型组合']
best_pred_data = model_predictions[best_model_name]
ax3.scatter(best_pred_data['y_true'], best_pred_data['y_pred'],
            c='darkred', alpha=0.7, s=50, edgecolors='white', linewidth=0.5)
ax3.plot([best_pred_data['y_true'].min(), best_pred_data['y_true'].max()],
         [best_pred_data['y_true'].min(), best_pred_data['y_true'].max()],
         'r--', lw=2, label='Ideal Prediction')
ax3.set_xlabel('True Concentration', fontsize=11)
ax3.set_ylabel('Predicted Concentration', fontsize=11)
ax3.set_title(f'Best Model: {best_model_name}\nAccuracy={best_pred_data["准确率"]}', fontsize=12, fontweight='bold')  # 替换R²为准确率
ax3.legend(fontsize=10)
ax3.grid(alpha=0.3)

# 5.4 预处理方法vs算法 热力图（修复索引问题）
algorithms = ['PLSR', 'XGBoost']
heatmap_data = np.zeros((6, 2))

for i, method in enumerate(preprocess_short_names):
    try:
        # PLSR准确率值
        pls_accuracy = metrics_df[(metrics_df['预处理方法'] == method) & (metrics_df['算法'] == 'PLSR')]['准确率'].values[0]
        # XGBoost准确率值
        xgb_accuracy = metrics_df[(metrics_df['预处理方法'] == method) & (metrics_df['算法'] == 'XGBoost')]['准确率'].values[0]
        heatmap_data[i, 0] = pls_accuracy
        heatmap_data[i, 1] = xgb_accuracy
    except:
        heatmap_data[i, 0] = 0
        heatmap_data[i, 1] = 0

# 绘制热力图
im = ax4.imshow(heatmap_data, cmap='RdYlGn', aspect='auto', vmin=0, vmax=1)
ax4.set_xticks(range(len(algorithms)))
ax4.set_yticks(range(len(preprocess_short_names)))
ax4.set_xticklabels(algorithms, fontsize=10)
ax4.set_yticklabels(preprocess_short_names, fontsize=10)
ax4.set_title('Accuracy Heatmap: Preprocessing Method vs. Algorithm', fontsize=12, fontweight='bold')  # 替换R²为准确率

# 添加数值标签
for i in range(len(preprocess_short_names)):
    for j in range(len(algorithms)):
        text = ax4.text(j, i, f'{heatmap_data[i, j]:.4f}',
                        ha="center", va="center", color="black", fontsize=9)

# 添加颜色条
cbar = plt.colorbar(im, ax=ax4, shrink=0.8)
cbar.set_label('Accuracy', fontsize=10)  # 替换R²值为准确率

plt.tight_layout()
plot_path = "O2-预处理方法+模型组合对比图.png"
plt.savefig(plot_path, dpi=300, bbox_inches='tight')
plt.close()
print(f"✅ 可视化图表已保存：{plot_path}")

# -------------------------- 6. 生成详细Word报告 --------------------------
print("\n=" * 60)
print("📌 生成详细Word分析报告")
doc = Document()

# 6.1 标题
title = doc.add_heading('超氧阴离子（O₂⁻）浓度预测 - 预处理方法×模型组合分析报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 6.2 数据说明
doc.add_heading('1. 实验设计说明', level=1)
doc.add_paragraph(f'• 光谱数据路径：{spectral_path}')
doc.add_paragraph(f'• 浓度数据路径：{concentration_path}')
doc.add_paragraph(f'• 实验方案：6种预处理方法 × 2种预测模型 = 12个组合模型')
doc.add_paragraph(f'• 预处理方法列表：')
for i, item in enumerate(preprocess_mapping):
    doc.add_paragraph(f'  - {item["short_name"]}：{item["sheet_name"]}')
doc.add_paragraph(f'• 预测模型：PLSR（偏最小二乘回归）、XGBoost（极端梯度提升）')
doc.add_paragraph(f'• 数据划分：训练集80%，测试集20%（随机种子=42）')
doc.add_paragraph(f'• 评估指标：准确率（决定系数）、MAE（平均绝对误差）、RMSE（均方根误差）、MAPE（平均绝对百分比误差）')  # 替换R²为准确率

# 6.3 所有模型评估结果
doc.add_heading('2. 所有模型评估结果', level=1)
# 按准确率排序的表格
table1 = doc.add_table(rows=len(metrics_df_sorted) + 1, cols=6)
table1.style = 'Table Grid'
# 表头
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = '排名'
hdr_cells[1].text = '模型组合'
hdr_cells[2].text = '准确率'  # 替换R²为准确率
hdr_cells[3].text = 'MAE'
hdr_cells[4].text = 'RMSE'
hdr_cells[5].text = 'MAPE（%）'

# 填充数据
for idx, row in metrics_df_sorted.iterrows():
    row_cells = table1.rows[idx + 1].cells
    row_cells[0].text = str(idx + 1)
    row_cells[1].text = row['模型组合']
    row_cells[2].text = str(row['准确率'])  # 替换R²为准确率
    row_cells[3].text = str(row['MAE'])
    row_cells[4].text = str(row['RMSE'])
    row_cells[5].text = str(row['MAPE'])

# 6.4 最优模型分析
doc.add_heading('3. 最优模型分析', level=1)
doc.add_paragraph(f'• 最优模型组合：{best_model["模型组合"]}')
doc.add_paragraph(f'• 核心指标：')
doc.add_paragraph(f'  - 准确率：{best_model["准确率"]}（越高越好）')  # 替换R²为准确率
doc.add_paragraph(f'  - MAE：{best_model["MAE"]}（越低越好）')
doc.add_paragraph(f'  - RMSE：{best_model["RMSE"]}（越低越好）')
doc.add_paragraph(f'  - MAPE：{best_model["MAPE"]}%（越低越好）')

# 分析最优模型的优势
preprocess_part = best_model['预处理方法']
algorithm_part = best_model['算法']
doc.add_paragraph(f'• 优势分析：')
doc.add_paragraph(f'  - 预处理方法优势：{preprocess_part}方法对光谱数据的预处理效果最佳')
doc.add_paragraph(f'  - 算法优势：{algorithm_part}算法更适合该预处理后的数据分布')

# 6.5 预处理方法对比分析
doc.add_heading('4. 预处理方法对比分析', level=1)
# 计算每种预处理方法的平均准确率
preprocess_avg = metrics_df.groupby('预处理方法')['准确率'].mean().sort_values(ascending=False)  # 替换R²为准确率
doc.add_paragraph('各预处理方法平均准确率（PLSR+XGBoost）：')
for method, avg_accuracy in preprocess_avg.items():
    doc.add_paragraph(f'• {method}：平均准确率 = {avg_accuracy:.4f}')  # 替换R²为准确率

# 6.6 算法对比分析
doc.add_heading('5. 算法对比分析', level=1)
algorithm_avg = metrics_df.groupby('算法')['准确率'].mean()  # 替换R²为准确率
doc.add_paragraph(f'• PLSR算法平均准确率：{algorithm_avg["PLSR"]:.4f}')  # 替换R²为准确率
doc.add_paragraph(f'• XGBoost算法平均准确率：{algorithm_avg["XGBoost"]:.4f}')  # 替换R²为准确率
doc.add_paragraph(f'• 结论：{"PLSR" if algorithm_avg["PLSR"] > algorithm_avg["XGBoost"] else "XGBoost"}算法整体表现更优')

# 6.7 可视化结果
doc.add_heading('6. 可视化对比结果', level=1)
doc.add_picture(plot_path, width=Inches(7))

# 6.8 结论与建议
doc.add_heading('7. 结论与建议', level=1)
conclusion = doc.add_paragraph()
conclusion.add_run(f'• 核心结论：{best_model["模型组合"]}是预测超氧阴离子浓度的最优组合，准确率达到{best_model["准确率"]}\n')  # 替换R²为准确率
conclusion.add_run(f'• 预处理方法效果排序：')
for i, (method, avg_accuracy) in enumerate(preprocess_avg.items(), 1):
    conclusion.add_run(f'{i}. {method}（平均准确率={avg_accuracy:.4f}）；')  # 替换R²为准确率
conclusion.add_run(f'\n• 算法选择建议：\n')
if algorithm_avg["PLSR"] > algorithm_avg["XGBoost"]:
    conclusion.add_run(f'  - 优先选择PLSR算法，整体稳定性和准确率更高；\n')
    conclusion.add_run(f'  - XGBoost算法可作为补充，针对特定预处理方法效果较好；\n')
else:
    conclusion.add_run(f'  - 优先选择XGBoost算法，非线性拟合能力更强；\n')
    conclusion.add_run(f'  - PLSR算法适合需要模型解释性的场景；\n')
conclusion.add_run(f'\n• 应用建议：\n')
conclusion.add_run(f'  - 实际应用中推荐使用{best_model["模型组合"]}模型；\n')
conclusion.add_run(f'  - 可进一步优化{preprocess_part}预处理方法的参数；\n')
conclusion.add_run(f'  - 可尝试将多种预处理方法的特征融合，进一步提升预测精度。')

# 保存报告
doc_path = "超氧阴离子浓度预测-预处理方法×模型组合详细分析报告.docx"
doc.save(doc_path)

# -------------------------- 7. 最终输出 --------------------------
print("\n=" * 60)
print("🎉 所有任务完成！")
print(f"📄 Word报告保存路径：{doc_path}")
print(f"📈 可视化图表保存路径：{plot_path}")
print(f"\n🏆 最优模型组合：{best_model['模型组合']}")
print(f"   准确率：{best_model['准确率']}")  # 替换R²为准确率
print(f"   均方根误差（RMSE）：{best_model['RMSE']}")
print("=" * 60)